# document_processor.py (without Celery)
import logging
from django.db import transaction
from .models import Document, DocumentChunk
from .pinecone_vector_store import PineconeVectorStore

logger = logging.getLogger(__name__)


# Same functions as before...
def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def read_pdf_file(file_path):
    try:
        import PyPDF2
        content = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        return content
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        raise

def read_docx_file(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise


def get_document_content(document):
    file_path = document.file.path
    file_type = document.file_type.lower()

    if file_type == 'txt':
        return read_text_file(file_path)
    elif file_type == 'pdf':
        return read_pdf_file(file_path)
    elif file_type in ['docx', 'doc']:
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def chunk_text(text, chunk_size=1000, overlap=200):
    """Splits text into overlapping chunks"""
    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)

        # Try to find a sentence break for clean chunking
        sentence_break = text.rfind('. ', start, end)
        if sentence_break != -1 and sentence_break > start + chunk_size // 2:
            end = sentence_break + 2

        chunks.append(text[start:end])
        start = max(start + chunk_size - overlap, end - overlap)

    return chunks


def extract_metadata(document):
    """Extract metadata from document based on file type"""
    metadata = {
        'filename': document.file.name,
        'file_type': document.file_type,
        'title': document.title,
        'document_id': str(document.id),
        'created_at': str(document.created_at),
        'updated_at': str(document.updated_at),
    }

    file_path = document.file.path
    file_type = document.file_type.lower()

    try:
        if file_type == 'pdf':
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if key.startswith('/'):
                            clean_key = key[1:].lower()
                            metadata[clean_key] = value
                metadata['page_count'] = len(reader.pages)

        elif file_type in ['docx', 'doc']:
            import docx
            doc = docx.Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)

            # Try to extract core properties
            try:
                core_props = doc.core_properties
                metadata['author'] = core_props.author
                metadata['created'] = str(core_props.created) if core_props.created else None
                metadata['modified'] = str(core_props.modified) if core_props.modified else None
            except:
                pass
    except Exception as e:
        logger.warning(f"Error extracting metadata from {document.title}: {e}")

    return metadata


# Removed @shared_task decorator
def process_document(document_id):
    try:
        document = Document.objects.get(id=document_id)
        document.status = 'processing'
        document.save(update_fields=['status'])

        # Get document content
        try:
            content = get_document_content(document)
        except Exception as e:
            logger.error(f"Error reading document content: {e}")
            document.status = 'failed'
            document.save(update_fields=['status'])
            return f"Failed to process document {document_id}: {str(e)}"

        # Extract metadata
        metadata = extract_metadata(document)

        # Add agent_id to metadata if available
        if document.agent:
            metadata['agent_id'] = str(document.agent.id)

        # Chunk the text
        chunks = chunk_text(content)
        total_chunks = len(chunks)

        if total_chunks == 0:
            document.status = 'failed'
            document.save(update_fields=['status'])
            return f"Failed to process document {document_id}: No content chunks generated"

        # Initialize Pinecone vector store
        vector_store = PineconeVectorStore()

        # Process chunks in batches
        try:
            # First create chunk records in the database
            chunk_objects = []
            for i, text_chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'chunk': i,
                    'total_chunks': total_chunks,
                    'character_start': i * (1000 - 200) if i > 0 else 0,
                    'character_end': min(i * (1000 - 200) + len(text_chunk), len(content))
                })

                chunk_obj = DocumentChunk(
                    document=document,
                    content=text_chunk,
                    chunk_index=i,
                    metadata=chunk_metadata
                )
                chunk_objects.append(chunk_obj)

            # Use bulk_create for efficiency
            with transaction.atomic():
                created_chunks = DocumentChunk.objects.bulk_create(chunk_objects)

            # Now process with Pinecone (generate embeddings and store vectors)
            chunk_ids = [chunk.id for chunk in created_chunks]
            chunk_texts = [chunk.content for chunk in created_chunks]
            chunk_metadatas = [chunk.metadata for chunk in created_chunks]

            # Add to user namespace
            embeddings = vector_store.add_documents(
                chunk_ids,
                chunk_texts,
                user_id=document.user.id,
                metadatas=chunk_metadatas
            )

            # Also add to global namespace
            vector_store.add_documents(
                chunk_ids,
                chunk_texts,
                metadatas=chunk_metadatas
            )

            # Update embeddings in database
            for i, embedding in enumerate(embeddings):
                created_chunks[i].embedding = embedding

            # Bulk update the embedding fields
            DocumentChunk.objects.bulk_update(created_chunks, ['embedding'])

            # Update document status
            document.total_chunks = total_chunks
            document.status = 'completed'
            document.save(update_fields=['total_chunks', 'status'])

            return f"Successfully processed document {document_id}: {total_chunks} chunks processed"

        except Exception as e:
            logger.error(f"Error processing chunks for document {document_id}: {e}")
            document.status = 'failed'
            document.save(update_fields=['status'])
            return f"Failed to process document {document_id}: {str(e)}"

    except Document.DoesNotExist:
        logger.error(f"Document with ID {document_id} not found")
        return f"Document with ID {document_id} not found"
    except Exception as e:
        logger.error(f"Unexpected error processing document {document_id}: {e}")
        try:
            document = Document.objects.get(id=document_id)
            document.status = 'failed'
            document.save(update_fields=['status'])
        except:
            pass
        return f"Failed to process document {document_id}: {str(e)}"